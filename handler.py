import json
import boto3
from io import BytesIO
from docx import Document  # To handle Word file generation

s3 = boto3.client('s3')

# Replace this with your S3 bucket name
BUCKET_NAME = 'foun-report'

def generate_pdf(event, context):
    try:
        # Parse input data from the event
        body = json.loads(event['body'])
        breaths_per_minute = body.get('breathsPerMinute')
        bmi = body.get('bmi')
        breath_hold_time = body.get('breathHoldTime')

        # Create a Word document
        doc = Document()
        doc.add_heading('Health Report', 0)

        # Add content based on the inputs that would normally be printed on the frontend
        if bmi is not None:
            if bmi < 18.5:
                doc.add_paragraph(f"Your BMI is {bmi}, which is considered underweight.")
            elif 18.5 <= bmi < 24.9:
                doc.add_paragraph(f"Your BMI is {bmi}, which is within the normal range.")
            elif 25 <= bmi < 29.9:
                doc.add_paragraph(f"Your BMI is {bmi}, which is considered overweight.")
            else:
                doc.add_paragraph(f"Your BMI is {bmi}, which is in the obese range.")

        if breaths_per_minute is not None:
            if breaths_per_minute < 12:
                doc.add_paragraph(f"You entered {breaths_per_minute} breaths per minute, which is below the normal range.")
            elif 12 <= breaths_per_minute <= 20:
                doc.add_paragraph(f"You entered {breaths_per_minute} breaths per minute, which is within the normal range.")
            else:
                doc.add_paragraph(f"You entered {breaths_per_minute} breaths per minute, which is above the normal range.")

        if breath_hold_time is not None:
            doc.add_paragraph(f"You held your breath for {breath_hold_time} seconds.")

        # Save Word file to a BytesIO stream
        word_file = BytesIO()
        doc.save(word_file)
        word_file.seek(0)

        # Generate a unique file name
        file_name = f"health_report_{bmi}_{breaths_per_minute}_{breath_hold_time}.docx"

        # Upload the Word document to S3
        s3.put_object(
            Bucket=BUCKET_NAME,
            Key=file_name,
            Body=word_file.getvalue(),
            ContentType='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

        # Generate a pre-signed URL for downloading the document
        download_url = s3.generate_presigned_url(
            'get_object',
            Params={'Bucket': BUCKET_NAME, 'Key': file_name},
            ExpiresIn=3600  # The URL will expire in 1 hour
        )

        # Return the download URL
        return {
            'statusCode': 200,
            'headers': {
                'Access-Control-Allow-Origin': '*',  # Allow all origins
                'Access-Control-Allow-Credentials': True  # Optional, needed if using credentials
            },
            'body': json.dumps({
                'message': 'Word document generated successfully',  # Updated to reflect Word file
                'downloadUrl': download_url  # Fixed variable name
            })
        }

    except Exception as e:
        return {
            'statusCode': 500,
            'body': json.dumps({'error': str(e)}),
            'headers': {'Content-Type': 'application/json'}
        }
